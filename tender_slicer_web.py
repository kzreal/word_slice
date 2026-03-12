#!/usr/bin/env python3
"""
标书切片工具 - Web 版本
使用浏览器访问 http://localhost:8000
"""

import io
import logging
import os
import zipfile
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

try:
    import requests
    import sys
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from dotenv import load_dotenv
    from flask import Flask, jsonify, render_template_string, request
    from werkzeug.exceptions import ClientDisconnected, RequestEntityTooLarge
except ImportError as e:
    print(f"缺少依赖: {e}")
    print("请运行: pip install -r requirements.txt")
    sys.exit(1)

# =============================================================================
# 日志配置
# =============================================================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('/tmp/tender_slicer_web.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)

# =============================================================================
# Flask 应用配置
# =============================================================================

app = Flask(__name__)

# Load environment variables
load_dotenv()

# LLM Configuration
LLM_API_ENDPOINT = os.getenv('LLM_API_ENDPOINT', '').strip()
LLM_API_KEY = os.getenv('LLM_API_KEY', '').strip()
LLM_MODEL = os.getenv('LLM_MODEL', 'vision-model').strip()
LLM_TIMEOUT = int(os.getenv('LLM_TIMEOUT', '30'))
LLM_MAX_RETRIES = int(os.getenv('LLM_MAX_RETRIES', '3'))

# Validate configuration
LLM_AVAILABLE = bool(LLM_API_ENDPOINT and LLM_API_KEY)
if LLM_AVAILABLE:
    logging.info("LLM image recognition enabled")
    logging.info(f"LLM endpoint: {LLM_API_ENDPOINT[:50]}...")
    logging.info(f"LLM model: {LLM_MODEL}")
else:
    logging.warning("LLM image recognition disabled - missing configuration")
    if not LLM_API_ENDPOINT:
        logging.warning("LLM_API_ENDPOINT not set")
    if not LLM_API_KEY:
        logging.warning("LLM_API_KEY not set")

UPLOAD_FOLDER = Path(__file__).parent / 'download'
UPLOAD_FOLDER.mkdir(exist_ok=True)

app.config.update({
    'UPLOAD_FOLDER': str(UPLOAD_FOLDER),
    'MAX_CONTENT_LENGTH': 1024 * 1024 * 1024,  # 1GB
})

logging.getLogger('werkzeug').setLevel(logging.ERROR)

# =============================================================================
# 标书切片器类
# =============================================================================

class ImageRecognitionService:
    """图像识别服务 - 通过 LLM API 识别图像内容"""

    def __init__(self, endpoint, api_key, model, timeout, max_retries):
        self.endpoint = endpoint
        self.api_key = api_key
        self.model = model
        self.timeout = timeout
        self.max_retries = max_retries

        # 创建 HTTP session
        self.session = requests.Session()
        self.session.headers.update({
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        })

        # 重试配置
        self.retry_adapter = requests.adapters.HTTPAdapter(
            max_retries=max_retries,
            pool_connections=10,
            pool_maxsize=100
        )
        self.session.mount('http://', self.retry_adapter)
        self.session.mount('https://', self.retry_adapter)

    def describe_image(self, image_data, image_format):
        """发送图像到 LLM API 获取描述"""
        # 重试逻辑
        for attempt in range(self.max_retries):
            try:
                # 编码图像为 base64
                import base64
                base64_image = base64.b64encode(image_data).decode('utf-8')

                # 构建请求体
                request_data = {
                    'model': self.model,
                    'messages': [
                        {
                            'role': 'user',
                            'content': [
                                {
                                    'type': 'text',
                                    'text': '十字以内直接概括图片的内容，不用加"图片展示..."、"这是...图片"等叙述'
                                },
                                {
                                    'type': 'image_url',
                                    'image_url': {
                                        'url': f'data:image/{image_format};base64,{base64_image}'
                                    }
                                }
                            ]
                        }
                    ],
                    'max_tokens': 100
                }

                # 发送请求
                response = self.session.post(
                    self.endpoint,
                    json=request_data,
                    timeout=self.timeout
                )

                response.raise_for_status()

                # 解析响应
                result = response.json()
                content = result['choices'][0]['message']['content'].strip()

                # 只返回纯文本描述，由调用者负责格式化
                return content

            except requests.exceptions.RequestException as e:
                logging.error(f"LLM API request failed (attempt {attempt + 1}/{self.max_retries}): {e}")
                if attempt == self.max_retries - 1:
                    # 重试次数用完，返回失败
                    logging.error(f"LLM API failed after {self.max_retries} attempts. Endpoint: {self.endpoint}")
                    return None
                # 等待一段时间再重试
                import time
                wait_time = 2 ** attempt
                logging.info(f"Waiting {wait_time}s before retry...")
                time.sleep(wait_time)  # 指数退避

            except (KeyError, IndexError) as e:
                logging.error(f"LLM API response parsing failed: {e}")
                logging.error(f"Response content: {response.text[:200]}...")
                return None

        return None

    def close(self):
        """关闭 session"""
        self.session.close()


class TenderSlicer:
    """标书切片器 - 将 Word 文档按目录大纲结构切片为多个编号的 Markdown 文件"""

    def __init__(self, docx_path):
        self.docx_path = Path(docx_path)
        self.doc = None
        self.sections = []

        # 初始化 LLM 服务
        if LLM_AVAILABLE:
            self.llm_service = ImageRecognitionService(
                endpoint=LLM_API_ENDPOINT,
                api_key=LLM_API_KEY,
                model=LLM_MODEL,
                timeout=LLM_TIMEOUT,
                max_retries=LLM_MAX_RETRIES
            )
        else:
            self.llm_service = None

    def iter_block_items(self, parent):
        """遍历文档中的所有块元素（段落和表格），保持原始顺序"""
        if hasattr(parent, 'element'):
            parent_elm = parent.element.body
        else:
            parent_elm = parent

        for element in parent_elm.iterchildren():
            if isinstance(element, CT_P):
                yield Paragraph(element, parent)
            elif isinstance(element, CT_Tbl):
                yield Table(element, parent)

    def get_image_format(self, image_element):
        """检测图片格式"""
        # 检查图片的 blip 元素
        blip = image_element.xpath('.//a:blip')
        if not blip:
            return 'png'  # 默认格式

        # 从 r:Embed 或 r:link 关系中获取图像
        embed = image_element.xpath('.//a:blip/@r:embed')
        if embed:
            # 这是嵌入的图像，从文档关系中获取
            try:
                image_part = self.doc.part.related_parts[embed[0]]
                return image_part.content_type.split('/')[-1]
            except KeyError:
                return 'png'

        return 'png'

    def get_image_from_relationship(self, image_ref):
        """从文档关系中提取实际的图像数据"""
        try:
            if hasattr(self.doc, 'part') and hasattr(self.doc.part, 'related_parts'):
                image_part = self.doc.part.related_parts[image_ref]
                return image_part.blob
            return None
        except Exception:
            return None

    def encode_image_to_base64(self, image_data, image_format):
        """将图像数据编码为 base64"""
        import base64
        return base64.b64encode(image_data).decode('utf-8')

    def extract_paragraph_images(self, paragraph):
        """检测段落中的图片，返回结构化图像数据"""
        images = []
        for run in paragraph.runs:
            for inline in run._element.xpath('.//w:drawing/wp:inline'):
                try:
                    blip = inline.xpath('.//a:blip')
                    if blip:
                        # 获取图像引用（使用关系 ID 作为唯一标识）
                        embed = inline.xpath('.//a:blip/@r:embed')
                        if embed:
                            # 提取实际图像数据
                            image_data = self.get_image_from_relationship(embed[0])
                            if image_data:
                                image_format = self.get_image_format(inline)
                                # 使用关系 ID 作为唯一标识
                                images.append({
                                    'id': embed[0],
                                    'data': image_data,
                                    'format': image_format,
                                    'placeholder': None  # 占位符由调用者根据行号生成
                                })
                except Exception as e:
                    logging.warning(f"Error extracting paragraph image: {e}")
        return images

    def extract_table_images(self, table):
        """提取表格单元格中的所有图片，返回结构化图像数据"""
        images = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    images.extend(self.extract_paragraph_images(paragraph))
        return images

    def load_document(self):
        """加载 Word 文档"""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.docx_path}")
        self.doc = Document(str(self.docx_path))

    def get_heading_level(self, paragraph):
        """获取段落的大纲级别，0=正文, 1=一级标题, 2=二级标题, ..."""
        # 方法1: 检查样式名称（支持 "Heading 1", "标题 1", "1. 标题" 等格式）
        if hasattr(paragraph, 'style') and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            # 支持中英文标题样式
            if 'heading 1' in style_name or '标题 1' in style_name or style_name == '1':
                return 1
            elif 'heading 2' in style_name or '标题 2' in style_name or style_name == '2':
                return 2
            elif 'heading 3' in style_name or '标题 3' in style_name or style_name == '3':
                return 3
            elif 'heading 4' in style_name or '标题 4' in style_name:
                return 4
            elif 'heading 5' in style_name or '标题 5' in style_name:
                return 5

        # 方法2: 检查大纲级别
        if hasattr(paragraph, '_element'):
            p = paragraph._element
            if p.pPr is not None and hasattr(p.pPr, 'outlineLvl') and p.pPr.outlineLvl is not None:
                return int(p.pPr.outlineLvl.val) + 1

        # 方法3: 通过字体大小和粗体判断（作为最后的备用方案）
        if hasattr(paragraph, 'runs') and paragraph.runs:
            try:
                first_run = paragraph.runs[0]
                if hasattr(first_run, 'bold') and first_run.bold:
                    if hasattr(first_run, 'font'):
                        if hasattr(first_run.font, 'size') and first_run.font.size:
                            size_pt = first_run.font.size.pt
                            # 根据字体大小推测标题级别
                            if size_pt >= 16:  # 16pt及以上可能是一级标题
                                return 1
                            elif size_pt >= 14:  # 14-16pt可能是二级标题
                                return 2
                            elif size_pt >= 12:  # 12-14pt可能是三级标题
                                return 3
            except Exception:
                pass

        return 0

    def table_to_markdown(self, table, start_no=1):
        """将表格转换为 Markdown 格式，带编号"""
        if not table.rows:
            return None, start_no

        lines = []
        no = start_no

        header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        lines.append(f"<!-- {no} --> | " + " | ".join(header_cells) + " |")
        lines.append("<!-- " + str(no + 1) + " --> |" + "|".join(["---"] * len(header_cells)) + "|")
        no += 2

        for row in table.rows[1:]:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if all(cell in ("", " ") for cell in cells):
                continue
            lines.append(f"<!-- {no} --> | " + " | ".join(cells) + " |")
            no += 1

        return '\n'.join(lines) + '\n\n', no

    def table_to_markdown_with_images(self, table, start_no, processed_images):
        """将表格转换为 Markdown 格式，支持图片合并"""
        if not table.rows:
            return None, start_no

        lines = []
        no = start_no

        # 收集表格中的所有图片
        table_images = []
        for row in table.rows:
            for cell in row.cells:
                table_images.extend(self.extract_table_images_for_row(cell))

        # 使用 LLM 处理表格图片（如果尚未处理）
        if table_images:
            for img in table_images:
                if img['id'] not in processed_images:
                    # 单个处理表格图片
                    if LLM_AVAILABLE and self.llm_service and img['data']:
                        try:
                            description = self.llm_service.describe_image(img['data'], img['format'])
                            if description:
                                processed_images[img['id']] = description
                            else:
                                processed_images[img['id']] = None
                        except Exception as e:
                            logging.error(f"Error processing table image {img['id']}: {e}")
                            processed_images[img['id']] = None
                    else:
                        processed_images[img['id']] = None

        # 处理表头
        header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        lines.append(f"<!-- {no} --> | " + " | ".join(header_cells) + " |")
        lines.append("<!-- " + str(no + 1) + " --> |" + "|".join(["---"] * len(header_cells)) + "|")
        no += 2

        # 处理数据行
        for row in table.rows[1:]:
            # 收集当前行的所有图片
            row_images = []
            for cell in row.cells:
                row_images.extend(self.extract_table_images_for_row(cell))

            # 如果该行有图片，合并输出
            if row_images:
                descriptions = []
                for img in row_images:
                    desc = processed_images.get(img['id'])
                    if desc:
                        descriptions.append(f"[图片: {desc}]")
                    else:
                        descriptions.append("[图片: 未识别图片]")

                # 合并图片作为新的一行，使用表格格式
                lines.append(f"<!-- {no} --> | " + " | ".join(descriptions) + " |")
                no += 1

            # 添加表格行内容
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not all(cell in ("", " ") for cell in cells):
                lines.append(f"<!-- {no} --> | " + " | ".join(cells) + " |")
                no += 1

        return '\n'.join(lines) + '\n\n', no

    def extract_table_images_for_row(self, cell):
        """从表格单元格中提取图片"""
        images = []
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                for inline in run._element.xpath('.//w:drawing/wp:inline'):
                    try:
                        blip = inline.xpath('.//a:blip')
                        if blip:
                            # 获取图像引用（使用关系 ID 作为唯一标识）
                            embed = inline.xpath('.//a:blip/@r:embed')
                            if embed:
                                # 提取实际图像数据
                                image_data = self.get_image_from_relationship(embed[0])
                                if image_data:
                                    image_format = self.get_image_format(inline)
                                    # 使用关系 ID 作为唯一标识
                                    images.append({
                                        'id': embed[0],
                                        'data': image_data,
                                        'format': image_format,
                                        'placeholder': None
                                    })
                    except Exception as e:
                        logging.warning(f"Error extracting table image: {e}")
        return images

    def sanitize_filename(self, filename):
        """清理文件名，移除非法字符"""
        import re
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        if len(filename) > 100:
            filename = filename[:100]
        return filename.strip()

    def process_images_with_llm(self, image_objects):
        """使用 LLM 处理所有图片，返回识别结果"""
        if not LLM_AVAILABLE or not self.llm_service:
            logging.warning("LLM service not available, returning None for all images")
            return {img['id']: None for img in image_objects}

        logging.info(f"Processing {len(image_objects)} images with LLM")

        results = {}
        for img in image_objects:
            try:
                description = self.llm_service.describe_image(
                    img['data'],
                    img['format']
                )
                if description:
                    # 只返回纯文本描述
                    results[img['id']] = description
                    logging.info(f"Image {img['id']} processed: {description[:50]}...")
                else:
                    # LLM 处理失败，返回 None
                    results[img['id']] = None
                    logging.warning(f"Failed to process image {img['id']}, returning None")
            except Exception as e:
                logging.error(f"Error processing image {img['id']}: {e}")
                results[img['id']] = None

        return results

    def process_images_batch(self, image_objects):
        """批量处理图片，优化性能"""
        if not LLM_AVAILABLE or not self.llm_service:
            logging.warning("LLM service not available, returning None for all images")
            return {img['id']: None for img in image_objects}

        logging.info(f"开始批量处理 {len(image_objects)} 张图片")

        # 过滤需要处理的图片
        images_to_process = []
        for img in image_objects:
            if img['data']:  # 只有有实际数据的图片才处理
                images_to_process.append(img)
            else:
                logging.warning(f"图片 {img['id']} 没有数据，跳过处理")

        if not images_to_process:
            logging.info("No images to process")
            return {img['id']: None for img in image_objects}

        # 限制图片大小（例如 10MB）
        processed_results = {}
        for i, img in enumerate(images_to_process, 1):
            try:
                logging.info(f"正在处理第 {i}/{len(images_to_process)} 张图片: {img['id']}")
                image_size = len(img['data'])
                logging.info(f"图片大小: {image_size / 1024:.2f} KB, 格式: {img['format']}")

                if image_size > 10 * 1024 * 1024:  # 10MB
                    processed_results[img['id']] = None
                    logging.warning(f"Image {img['id']} too large, skipping")
                    continue

                logging.info(f"开始调用 LLM API 识别图片 {img['id']}")
                description = self.llm_service.describe_image(img['data'], img['format'])
                logging.info(f"LLM API 调用完成，结果长度: {len(description) if description else 0}")

                if description:
                    # 只返回纯文本描述
                    processed_results[img['id']] = description
                    logging.info(f"图片 {img['id']} 处理成功")
                else:
                    processed_results[img['id']] = None
                    logging.warning(f"图片 {img['id']} 识别失败")
            except Exception as e:
                logging.error(f"Error processing image {img['id']}: {e}")
                logging.exception(f"图片 {img['id']} 处理异常详情")
                processed_results[img['id']] = None

        # 为未处理的图片添加 None
        for img in image_objects:
            if img['id'] not in processed_results:
                processed_results[img['id']] = None

        return processed_results

    def cleanup(self):
        """清理资源"""
        if self.llm_service:
            self.llm_service.close()

    def slice_document(self, max_level=None):
        """切片文档，每行添加编号，保留所有表格和图片

        max_level: 切片级别
            0 - 零级模式：不按章节切片，整个文档转换为一个 Markdown 文件
            1, 2, 3 - 按指定级别标题切片
            None 或 'all' - 按所有标题层级切片
        """
        self.load_document()

        # 收集所有图片
        all_images = []

        # 遍历文档收集图片
        logging.info(f"开始提取图片，文档包含 {len(self.doc.paragraphs)} 个段落和 {len(self.doc.tables)} 个表格")
        for paragraph in self.doc.paragraphs:
            images = self.extract_paragraph_images(paragraph)
            all_images.extend(images)

        for table in self.doc.tables:
            # 为了避免重复，这里不需要再次收集表格图片
            # 表格图片会在 table_to_markdown_with_images 中处理
            pass

        logging.info(f"共提取到 {len(all_images)} 张图片")

        # 使用 LLM 处理图片
        processed_images = {}
        if LLM_AVAILABLE and self.llm_service and all_images:
            logging.info(f"开始使用 LLM 处理 {len(all_images)} 张图片")
            # 使用批处理提高性能
            processed_images = self.process_images_batch(all_images)
            logging.info(f"LLM 图片处理完成，成功处理 {len([k for k, v in processed_images.items() if v])} 张图片")
        else:
            # LLM 不可用时设置为 None，由调用者生成占位符
            logging.warning(f"LLM 不可用或没有图片")
            for img in all_images:
                processed_images[img['id']] = None

        # 零级模式：不按章节切片，整个文档为一个 Markdown 文件
        if max_level == 0:
            sections = []
            line_no = 1

            full_section = {
                'level': 0,
                'title': self.docx_path.stem,
                'content': [],
                'index': 0
            }

            for block in self.iter_block_items(self.doc):
                if isinstance(block, Paragraph):
                    level = self.get_heading_level(block)
                    text = block.text.strip()
                    images = self.extract_paragraph_images(block)

                    if not text and not images:
                        continue

                    if any(kw in text for kw in ['目录', '目  录', 'CONTENTS']):
                        continue

                    # 保留原始标题结构
                    if level > 0:
                        full_section['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                        line_no += 1
                    elif text:
                        full_section['content'].append(f"<!-- {line_no} --> {text}\n")
                        line_no += 1

                    for img in images:
                        description = processed_images.get(img['id'])
                        if description:
                            # LLM 识别结果，分离行号和图片
                            full_section['content'].append(f"<!-- {line_no} -->[图片: {description}]\n")
                        else:
                            # 使用占位符，分离行号和图片
                            full_section['content'].append(f"<!-- {line_no} -->[图片: 未识别图片]\n")
                        line_no += 1

                elif isinstance(block, Table):
                    table_md, line_no = self.table_to_markdown_with_images(block, line_no, processed_images)
                    if table_md:
                        full_section['content'].append(table_md)

            if full_section['content']:
                sections.append(full_section)

            self.sections = sections
            return sections

        # 原有逻辑：按章节层级切片
        if max_level is None or max_level == 'all':
            max_level = float('inf')

        sections = []
        section_stack = []
        section_index = 0
        line_no = 1

        cover_section = {
            'level': 0,
            'title': '封面',
            'content': [],
            'index': section_index
        }
        section_stack.append(cover_section)

        for block in self.iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                level = self.get_heading_level(block)
                text = block.text.strip()
                images = self.extract_paragraph_images(block)

                if not text and not images:
                    continue

                if any(kw in text for kw in ['目录', '目  录', 'CONTENTS']):
                    continue

                if level > 0:
                    if level <= max_level:
                        if section_stack[-1]['content']:
                            sections.append(section_stack[-1])
                            section_index += 1

                        new_section = {
                            'level': level,
                            'title': text,
                            'content': [],
                            'index': section_index
                        }

                        while section_stack and section_stack[-1]['level'] > level:
                            section_stack.pop()

                        section_stack.append(new_section)
                        section_stack[-1]['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                        line_no += 1
                    else:
                        section_stack[-1]['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                        line_no += 1
                else:
                    if text:
                        section_stack[-1]['content'].append(f"<!-- {line_no} --> {text}\n")
                        line_no += 1

                for img in images:
                    description = processed_images.get(img['id'])
                    if description:
                        # LLM 识别结果，分离行号和图片
                        section_stack[-1]['content'].append(f"<!-- {line_no} -->[图片: {description}]\n")
                    else:
                        # 使用占位符，分离行号和图片
                        section_stack[-1]['content'].append(f"<!-- {line_no} -->[图片: 未识别图片]\n")
                    line_no += 1

            elif isinstance(block, Table):
                table_md, line_no = self.table_to_markdown_with_images(block, line_no, processed_images)
                if table_md:
                    section_stack[-1]['content'].append(table_md)

        for section in section_stack:
            if section['content']:
                sections.append(section)

        self.sections = sections
        return sections


# =============================================================================
# 错误处理
# =============================================================================

@app.errorhandler(RequestEntityTooLarge)
def handle_request_entity_too_large(e):
    return jsonify({'error': '文件大小超过限制，最大支持 1 GB'}), 413


@app.errorhandler(ClientDisconnected)
def handle_client_disconnected(e):
    return jsonify({'error': '上传中断'}), 400


# =============================================================================
# 路由定义
# =============================================================================

@app.route('/')
def index():
    return render_template_string(_INDEX_HTML)


@app.route('/slice', methods=['POST'])
def slice_file():
    """切片文件接口，返回 zip 文件"""
    upload_path = None
    logging.info("=" * 50)
    logging.info("收到新的切片请求")
    try:
        if 'file' not in request.files:
            logging.error("请求中没有文件")
            return jsonify({'error': '未上传文件'}), 400

        file = request.files['file']
        logging.info(f"文件名: {file.filename}, 文件大小: {file.content_length / 1024:.2f} KB")

        if file.filename == '':
            logging.error("文件名为空")
            return jsonify({'error': '文件名为空'}), 400

        if not file.filename.endswith('.docx'):
            logging.error(f"不支持的文件格式: {file.filename}")
            return jsonify({'error': '只支持 .docx 格式'}), 400

        max_level = request.form.get('max_level', '0')
        logging.info(f"切片级别: {max_level}")

        try:
            if max_level.lower() == 'all':
                max_level = None  # 全部层级
            else:
                max_level = int(max_level)
                if max_level not in [0, 1, 2, 3]:
                    logging.error(f"无效的切分层级: {max_level}")
                    return jsonify({'error': '无效的切分层级'}), 400
        except ValueError as e:
            logging.error(f"切分层级参数格式错误: {e}")
            return jsonify({'error': '切分层级参数格式错误'}), 400

        upload_path = UPLOAD_FOLDER / file.filename
        logging.info(f"开始保存文件到: {upload_path}")
        file.save(str(upload_path))
        logging.info(f"文件保存成功")

        # 验证文件
        if upload_path.stat().st_size == 0:
            logging.error(f"文件大小为 0，可能是上传失败")
            upload_path.unlink()
            return jsonify({'error': '上传的文件为空，请重新上传'}), 400

        # 检查是否为有效的 ZIP 文件（docx 本质是 ZIP）
        try:
            import zipfile
            with zipfile.ZipFile(upload_path, 'r') as zf:
                # 检查必要的文件是否存在
                required_files = ['[Content_Types].xml', '_rels/.rels', 'word/document.xml']
                for req_file in required_files:
                    try:
                        zf.getinfo(req_file)
                    except KeyError:
                        logging.error(f"docx 文件缺少必要文件: {req_file}")
                        upload_path.unlink()
                        return jsonify({'error': f'docx 文件格式不正确，缺少 {req_file}'}), 400
        except zipfile.BadZipFile:
            logging.error("文件不是有效的 ZIP 文件")
            upload_path.unlink()
            return jsonify({'error': '文件不是有效的 docx 格式'}), 400
        except Exception as e:
            logging.error(f"验证文件时出错: {e}")
            upload_path.unlink()
            return jsonify({'error': f'文件验证失败: {str(e)}'}), 400

        logging.info("开始切片文档")
        slicer = TenderSlicer(upload_path)
        try:
            sections = slicer.slice_document(max_level=max_level)
            logging.info(f"切片完成，共生成 {len(sections)} 个切片")
        except KeyError as e:
            error_msg = str(e)
            if "NULL" in error_msg or "archive" in error_msg:
                # 特殊处理 docx 文件损坏的情况
                logging.error("docx 文件内部结构损坏或格式异常")
                slicer.cleanup()
                if upload_path.exists():
                    upload_path.unlink()
                return jsonify({'error': 'docx 文件内部结构损坏或格式异常。请尝试用 Microsoft Word 打开文件并另存为新的 docx 文件后再上传。'}), 400
            raise
        finally:
            slicer.cleanup()

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            logging.info(f"开始生成 ZIP 文件，包含 {len(sections)} 个章节")
            for i, section in enumerate(sections, 1):
                index = section['index'] + 1
                index_str = str(index).zfill(3)
                safe_title = slicer.sanitize_filename(section['title'])
                filename = f"{index_str}_{safe_title}.md"
                content = ''.join(section['content'])
                zipf.writestr(filename, content.encode('utf-8'))
                logging.debug(f"已添加第 {i}/{len(sections)} 个文件: {filename}")

            index_content = "# 标书切片索引\n\n"
            index_content += f"原文件: {file.filename}\n"
            index_content += f"切片时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            index_content += f"总章节数: {len(sections)}\n\n---\n\n## 章节列表\n\n"

            for section in sections:
                idx = section['index'] + 1
                title = section['title']
                filename = f"{str(idx).zfill(3)}_{slicer.sanitize_filename(title)}.md"
                index_content += f"- [{idx}. {title}]({filename})\n"

            zipf.writestr("00_index.md", index_content.encode('utf-8'))

        zip_buffer.seek(0)
        logging.info(f"ZIP 文件生成完成，大小: {zip_buffer.getbuffer().nbytes / 1024:.2f} KB")

        try:
            upload_path.unlink()
        except Exception:
            pass

        safe_filename = quote(f"sliced_{file.filename}.zip", safe='')
        logging.info(f"准备发送响应，文件名: {safe_filename}")
        logging.info("切片请求处理成功完成")
        logging.info("=" * 50)
        return app.response_class(
            zip_buffer.getvalue(),
            mimetype='application/zip',
            headers={
                'Content-Disposition': f"attachment; filename*=UTF-8''{safe_filename}",
                'X-Section-Count': str(len(sections))
            }
        )

    except Exception as e:
        import traceback
        logging.error("切片请求处理失败")
        logging.error(f"错误类型: {type(e).__name__}")
        logging.error(f"错误信息: {str(e)}")
        logging.error(f"堆栈跟踪:\n{traceback.format_exc()}")

        if upload_path and upload_path.exists():
            try:
                upload_path.unlink()
            except Exception:
                pass

        error_msg = str(e)
        if len(error_msg) > 500:
            error_msg = error_msg[:500] + '...'
        logging.error(f"返回错误响应: {error_msg}")
        logging.info("=" * 50)
        return jsonify({'error': error_msg}), 500


# =============================================================================
# HTML 模板
# =============================================================================

_INDEX_HTML = '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>标书切片工具</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 20px;
        }
        .container {
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            padding: 40px;
            max-width: 600px;
            width: 100%;
        }
        h1 { color: #333; margin-bottom: 30px; text-align: center; font-size: 28px; }
        .upload-area {
            border: 3px dashed #667eea;
            border-radius: 15px;
            padding: 50px 20px;
            text-align: center;
            background: #f8f9ff;
            transition: all 0.3s;
            cursor: pointer;
        }
        .upload-area:hover, .upload-area.dragover {
            border-color: #764ba2;
            background: #f0f2ff;
        }
        .icon { font-size: 48px; margin-bottom: 15px; }
        .text { color: #666; font-size: 16px; }
        .subtext { color: #999; font-size: 14px; margin-top: 8px; }
        #fileInput { display: none; }
        .btn {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 12px 30px;
            border-radius: 25px;
            font-size: 16px;
            cursor: pointer;
            display: block;
            margin: 20px auto 0;
        }
        .btn:hover { transform: translateY(-2px); box-shadow: 0 5px 20px rgba(102, 126, 234, 0.4); }
        .btn:disabled { background: #ccc; cursor: not-allowed; transform: none; }
        #fileInfo {
            background: #f0f2ff;
            padding: 15px;
            border-radius: 10px;
            margin-top: 20px;
            display: none;
        }
        #fileInfo.show { display: block; }
        #progress {
            margin-top: 20px;
            display: none;
        }
        #progress.show { display: block; }
        .progress-bar {
            background: #e0e0e0;
            border-radius: 10px;
            height: 8px;
            overflow: hidden;
        }
        .progress-fill {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            height: 100%;
            width: 0%;
            transition: width 0.3s;
        }
        #result { margin-top: 20px; display: none; }
        #result.show { display: block; }
        .success {
            background: #f0fff4;
            border: 1px solid #52C41A;
            color: #389e0d;
            padding: 15px;
            border-radius: 10px;
            text-align: center;
        }
        .download-btn {
            background: #52C41A;
            color: white;
            border: none;
            padding: 14px 28px;
            border-radius: 30px;
            font-size: 16px;
            cursor: pointer;
            margin-top: 10px;
            text-decoration: none;
            display: inline-block;
        }
        .download-btn:hover { background: #389e0d; }
        .error {
            background: #fff2f0;
            border: 1px solid #ff4d4f;
            color: #ff4d4f;
            padding: 15px;
            border-radius: 10px;
        }
        #levelSelector {
            margin-top: 20px;
            background: #f8f9ff;
            padding: 20px;
            border-radius: 12px;
        }
        #levelSelector label {
            display: block;
            color: #666;
            font-size: 14px;
            font-weight: 500;
            margin-bottom: 12px;
        }
        .level-options { display: flex; gap: 10px; }
        .level-option { cursor: pointer; position: relative; }
        .level-option input[type="radio"] { display: none; }
        .level-option span {
            display: inline-block;
            padding: 10px 20px;
            background: white;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            color: #666;
            font-size: 14px;
            transition: all 0.3s;
        }
        .level-option input[type="radio"]:checked + span {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border-color: #667eea;
            color: white;
        }
        .level-option:hover span { border-color: #667eea; }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 标书切片工具</h1>

        <div class="upload-area" id="uploadArea">
            <div class="icon">📁</div>
            <div class="text">点击或拖拽上传投标文件</div>
            <div class="subtext">支持 .docx 格式，最大 1GB</div>
        </div>

        <input type="file" id="fileInput" accept=".docx">

        <div id="fileInfo"></div>

        <div id="levelSelector">
            <label>切分层级</label>
            <div class="level-options">
                <label class="level-option">
                    <input type="radio" name="sliceLevel" value="0">
                    <span>零级</span>
                </label>
                <label class="level-option">
                    <input type="radio" name="sliceLevel" value="1">
                    <span>一级</span>
                </label>
                <label class="level-option">
                    <input type="radio" name="sliceLevel" value="2" checked>
                    <span>二级</span>
                </label>
                <label class="level-option">
                    <input type="radio" name="sliceLevel" value="3">
                    <span>三级</span>
                </label>
                <label class="level-option">
                    <input type="radio" name="sliceLevel" value="all">
                    <span>全部</span>
                </label>
            </div>
        </div>

        <button class="btn" id="sliceBtn" disabled>开始切片</button>

        <div id="progress">
            <div style="text-align: center; color: #666; margin-bottom: 8px;" id="progressText">准备中...</div>
            <div class="progress-bar">
                <div class="progress-fill" id="progressFill"></div>
            </div>
        </div>

        <div id="result"></div>
    </div>

    <script>
        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        const fileInfo = document.getElementById('fileInfo');
        const sliceBtn = document.getElementById('sliceBtn');
        const progress = document.getElementById('progress');
        const progressFill = document.getElementById('progressFill');
        const progressText = document.getElementById('progressText');
        const result = document.getElementById('result');

        let selectedFile = null;

        uploadArea.addEventListener('click', () => fileInput.click());

        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });

        uploadArea.addEventListener('dragleave', () => {
            uploadArea.classList.remove('dragover');
        });

        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            if (e.dataTransfer.files.length > 0) {
                handleFile(e.dataTransfer.files[0]);
            }
        });

        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length > 0) {
                handleFile(e.target.files[0]);
            }
        });

        function handleFile(file) {
            const MAX_SIZE = 1024 * 1024 * 1024;
            if (file.size > MAX_SIZE) {
                showError('文件大小超过 1GB 限制');
                return;
            }

            if (!file.name.endsWith('.docx')) {
                showError('请上传 .docx 格式的文件');
                return;
            }

            selectedFile = file;
            fileInfo.innerHTML = `✓ 已选择: <strong>${file.name}</strong><br>大小: ${formatFileSize(file.size)}`;
            fileInfo.classList.add('show');
            sliceBtn.disabled = false;
            result.innerHTML = '';
            result.classList.remove('show');
        }

        function formatFileSize(bytes) {
            if (bytes < 1024) return bytes + ' B';
            if (bytes < 1024 * 1024) return (bytes / 1024).toFixed(2) + ' KB';
            if (bytes < 1024 * 1024 * 1024) return (bytes / (1024 * 1024)).toFixed(2) + ' MB';
            return (bytes / (1024 * 1024 * 1024)).toFixed(2) + ' GB';
        }

        function showError(message) {
            result.innerHTML = `<div class="error">${message}</div>`;
            result.classList.add('show');
            sliceBtn.disabled = true;
            fileInfo.classList.remove('show');
        }

        sliceBtn.addEventListener('click', async () => {
            if (!selectedFile) return;

            progress.classList.add('show');
            result.innerHTML = '';
            result.classList.remove('show');
            sliceBtn.disabled = true;

            const formData = new FormData();
            formData.append('file', selectedFile);
            formData.append('max_level', document.querySelector('input[name="sliceLevel"]:checked').value);

            const startTime = Date.now();

            // 使用 XMLHttpRequest 实现真实的上传进度
            return new Promise((resolve, reject) => {
                const xhr = new XMLHttpRequest();

                // 上传进度
                xhr.upload.addEventListener('progress', (e) => {
                    if (e.lengthComputable) {
                        const percent = Math.min(90, Math.round((e.loaded / e.total) * 80));
                        const uploadedMB = (e.loaded / 1024 / 1024).toFixed(2);
                        const totalMB = (e.total / 1024 / 1024).toFixed(2);
                        const elapsed = (Date.now() - startTime) / 1000;
                        const speedMBps = (e.loaded / 1024 / 1024 / elapsed).toFixed(2);
                        updateProgress(percent, `上传中... ${uploadedMB}MB / ${totalMB}MB (${speedMBps}MB/s)`);
                    }
                });

                xhr.addEventListener('load', () => {
                    if (xhr.status === 200) {
                        updateProgress(90, '正在生成切片...');
                        const blob = new Blob([xhr.response], { type: 'application/zip' });
                        const url = URL.createObjectURL(blob);

                        const sectionCount = xhr.getResponseHeader('X-Section-Count') || '多个';
                        result.innerHTML = `<div class="success">✅ 切片完成！共 ${sectionCount} 个章节</div>`;

                        const downloadBtn = document.createElement('a');
                        downloadBtn.href = url;
                        downloadBtn.download = 'sliced_documents.zip';
                        downloadBtn.className = 'download-btn';
                        downloadBtn.textContent = '⬇️ 下载切片结果';
                        downloadBtn.style.display = 'block';
                        downloadBtn.style.textAlign = 'center';
                        result.appendChild(downloadBtn);
                        result.classList.add('show');

                        updateProgress(100, '处理完成！');
                        sliceBtn.disabled = false;
                        resolve();
                    } else {
                        let errorMsg = '处理失败';
                        try {
                            const errorData = JSON.parse(xhr.responseText);
                            errorMsg = errorData.error || errorMsg;
                        } catch {}
                        showError(`❌ ${errorMsg}`);
                        sliceBtn.disabled = false;
                        reject(new Error(errorMsg));
                    }
                });

                xhr.addEventListener('error', () => {
                    showError('❌ 网络错误，请重试');
                    sliceBtn.disabled = false;
                    reject(new Error('网络错误'));
                });

                xhr.addEventListener('timeout', () => {
                    showError('❌ 请求超时，请重试');
                    sliceBtn.disabled = false;
                    reject(new Error('请求超时'));
                });

                xhr.responseType = 'blob';
                xhr.timeout = Math.max(3600000, selectedFile.size * 0.005);
                xhr.open('POST', '/slice');
                xhr.send(formData);
            });
        });

        function updateProgress(value, text) {
            progressFill.style.width = value + '%';
            progressText.textContent = text;
        }
    </script>
</body>
</html>'''


# =============================================================================
# 主程序入口
# =============================================================================

if __name__ == '__main__':
    print("标书切片工具 - Web 版本")
    print("请在浏览器中访问: http://localhost:8000")
    print("局域网访问: http://[您的IP地址]:8000")
    print("按 Ctrl+C 停止服务器")
    app.run(host='0.0.0.0', port=8000, debug=True, threaded=True)
