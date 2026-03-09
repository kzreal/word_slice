#!/usr/bin/env python3
"""
标书文档切片脚本
将 Word (.docx) 文档按目录大纲结构切片为多个编号的 Markdown 文件
"""

import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


class TenderSlicer:
    """标书切片器 - 将 Word 文档按目录大纲结构切片为多个编号的 Markdown 文件"""

    def __init__(self, docx_path: str, output_dir: str = "sliced", slice_level: int = None):
        """
        初始化切片器

        Args:
            docx_path: 输入的 Word 文档路径
            output_dir: 输出目录，默认为 "sliced"
            slice_level: 切片级别，0 表示不按章节切片（整个文档为一个 md 文件），
                        None 表示按所有标题层级切片，1 表示只按一级标题切片，以此类推
        """
        self.docx_path = Path(docx_path)
        self.output_dir = Path(output_dir)
        self.doc = None
        self.sections = []  # 存储切片结果
        self.slice_level = slice_level  # 切片级别

        # 创建输出目录
        self.output_dir.mkdir(exist_ok=True)

    def load_document(self):
        """加载 Word 文档"""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.docx_path}")
        self.doc = Document(str(self.docx_path))

    def get_heading_level(self, paragraph):
        """
        获取段落的标题级别

        Args:
            paragraph: 段落对象

        Returns:
            int: 0=正文, 1=一级标题, 2=二级标题, ...
        """
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            # 提取标题级别，如 "Heading 1" -> 1
            match = re.search(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
            if match:
                return int(match.group(1))

        # 检查大纲级别（有些文档使用大纲级别而不是标题样式）
        if hasattr(paragraph, '_element'):
            p = paragraph._element
            if p.pPr and p.pPr.outlineLvl is not None:
                # outlineLvl 是从0开始的，所以加1
                return int(p.pPr.outlineLvl.val) + 1
        return 0

    def get_paragraph_text(self, paragraph):
        """
        获取段落文本

        Args:
            paragraph: 段落对象

        Returns:
            str or None: 段落文本，如果为空则返回 None
        """
        text = paragraph.text.strip()
        if text:
            return text
        return None

    def convert_to_markdown(self):
        """
        将整个 Word 文档转换成 Markdown 格式

        Returns:
            str: 完整的 Markdown 内容
        """
        elements = []

        # 收集段落及其位置
        for paragraph in self.doc.paragraphs:
            position = paragraph._element.getparent().index(paragraph._element)
            elements.append({
                'type': 'paragraph',
                'position': position,
                'data': paragraph
            })

        # 收集表格及其位置
        for table in self.doc.tables:
            position = table._element.getparent().index(table._element)
            elements.append({
                'type': 'table',
                'position': position,
                'data': table
            })

        # 按位置排序
        elements.sort(key=lambda x: x['position'])

        # 按顺序转换成 Markdown
        md_lines = []
        for element in elements:
            if element['type'] == 'paragraph':
                paragraph = element['data']
                level = self.get_heading_level(paragraph)
                text = self.get_paragraph_text(paragraph)

                # 跳过空段落
                if not text:
                    continue

                # 检测目录页（跳过）
                if any(keyword in text for keyword in ['目录', '目  录', 'CONTENTS', 'TABLE OF CONTENTS']):
                    continue

                # 保留原始标题结构
                if level > 0:
                    md_lines.append(f"{'#' * level} {text}\n")
                else:
                    md_lines.append(f"{text}\n")

            elif element['type'] == 'table':
                table_md = self.table_to_markdown(element['data'])
                if table_md:
                    md_lines.append(table_md)

        return ''.join(md_lines)

    def _parse_markdown_structure(self, md_content):
        """
        解析 Markdown 内容，提取标题和对应的内容块

        Args:
            md_content: Markdown 字符串

        Returns:
            list: 标题信息列表，每个元素包含 {level, title, start_line}
        """
        lines = md_content.split('\n')
        headings = []

        for i, line in enumerate(lines):
            if line.startswith('#'):
                # 计算标题级别
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break

                # 提取标题文本（去掉 # 和空格）
                title = line[level:].strip()
                if title:
                    headings.append({
                        'level': level,
                        'title': title,
                        'start_line': i
                    })

        return headings, lines

    def _slice_markdown_by_level(self, md_content, slice_level):
        """
        根据 Markdown 内容按标题级别进行切片

        Args:
            md_content: Markdown 字符串
            slice_level: 切片级别（0=零级，None=全部，1=一级，2=二级，3=三级）

        Returns:
            list: 切片后的章节列表
        """
        headings, lines = self._parse_markdown_structure(md_content)

        if slice_level == 0:
            # 零级模式：不切片
            return [{
                'level': 0,
                'title': self.docx_path.stem,
                'content': md_content,
                'index': 0
            }]

        # 构建章节
        sections = []
        section_index = 0

        # 添加封面章节（第一个标题之前的内容）
        if headings:
            cover_content = '\n'.join(lines[:headings[0]['start_line']])
            if cover_content.strip():
                sections.append({
                    'level': 0,
                    'title': '封面',
                    'content': cover_content + '\n\n',
                    'index': section_index
                })
                section_index += 1

        # 确定最大切片级别
        max_level = float('inf') if slice_level is None else slice_level

        # 根据标题级别构建章节
        for i, heading in enumerate(headings):
            level = heading['level']
            title = heading['title']

            # 确定这个标题应该属于哪个章节
            # 如果标题级别 <= max_level，则创建新章节
            # 否则，这个标题应该属于上一个章节

            if level <= max_level:
                # 确定章节内容范围
                start_line = heading['start_line']
                if i + 1 < len(headings):
                    # 找到下一个级别 <= max_level 的标题
                    end_line = headings[i + 1]['start_line']
                    for j in range(i + 1, len(headings)):
                        if headings[j]['level'] <= max_level:
                            end_line = headings[j]['start_line']
                            break
                else:
                    end_line = len(lines)

                # 提取章节内容
                section_content = '\n'.join(lines[start_line:end_line]) + '\n\n'

                sections.append({
                    'level': level,
                    'title': title,
                    'content': section_content,
                    'index': section_index
                })
                section_index += 1

        return sections

    def slice_document(self):
        """
        切片文档，按照标题级别进行切片

        流程：
        1. 将 Word 文档完整转换成 Markdown
        2. 根据 slice_level 对 Markdown 进行切片
        """
        if not self.doc:
            self.load_document()

        # 第一步：将 Word 文档转换成 Markdown
        md_content = self.convert_to_markdown()

        # 第二步：根据切片级别对 Markdown 进行切片
        self.sections = self._slice_markdown_by_level(md_content, self.slice_level)

    def table_to_markdown(self, table):
        """
        将表格转换为 Markdown 格式

        Args:
            table: 表格对象

        Returns:
            str or None: Markdown 格式的表格，如果表格为空则返回 None
        """
        if not table.rows:
            return None

        markdown = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            markdown.append('| ' + ' | '.join(cells) + ' |')

            # 添加分隔线（第一行之后）
            if i == 0:
                separator = '|' + '|'.join(['---'] * len(cells)) + '|'
                markdown.append(separator)

        return '\n'.join(markdown) + '\n\n'

    def sanitize_filename(self, filename):
        """
        清理文件名，移除非法字符

        Args:
            filename: 原始文件名

        Returns:
            str: 清理后的安全文件名
        """
        # 移除或替换非法字符
        illegal_chars = r'[<>:"/\\|?*]'
        filename = re.sub(illegal_chars, '', filename)
        # 限制长度
        if len(filename) > 100:
            filename = filename[:100]
        return filename.strip()

    def save_sections(self):
        """保存切片后的文件"""
        print(f"正在保存 {len(self.sections)} 个章节...")

        for section in self.sections:
            # 生成编号
            index = section['index'] + 1
            index_str = str(index).zfill(3)  # 补零到3位

            # 清理标题作为文件名
            safe_title = self.sanitize_filename(section['title'])
            filename = f"{index_str}_{safe_title}.md"
            filepath = self.output_dir / filename

            # 写入内容
            content = ''.join(section['content'])
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)

            print(f"  已保存: {filename}")

        print(f"\n完成！所有文件已保存到: {self.output_dir}")

    def generate_index(self):
        """生成索引文件"""
        index_file = self.output_dir / "00_index.md"
        with open(index_file, 'w', encoding='utf-8') as f:
            f.write("# 标书切片索引\n\n")
            f.write(f"原文件: {self.docx_path.name}\n")
            f.write(f"切片时间: {self._get_timestamp()}\n")
            f.write(f"总章节数: {len(self.sections)}\n")
            if self.slice_level is not None:
                f.write(f"切片级别: {self.slice_level} ({'零级模式-整个文档为一个文件' if self.slice_level == 0 else f'按 {self.slice_level} 级标题切片'})\n")
            f.write("\n---\n\n")
            f.write("## 章节列表\n\n")

            for section in self.sections:
                index = section['index'] + 1
                level = section['level']
                title = section['title']
                filename = f"{str(index).zfill(3)}_{self.sanitize_filename(title)}.md"

                # 根据级别添加缩进
                indent = "  " * (level - 1)
                f.write(f"{indent}- [{index}. {title}]({filename})\n")

        print(f"索引文件已生成: {index_file}")

    @staticmethod
    def _get_timestamp():
        """获取当前时间戳"""
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def process(self):
        """执行完整的切片流程"""
        print(f"开始处理: {self.docx_path}")
        print("-" * 50)
        self.load_document()
        self.slice_document()
        self.save_sections()
        self.generate_index()
        print("-" * 50)
        print("处理完成！")


def main():
    """主函数"""
    import sys

    if len(sys.argv) < 2:
        print("使用方法: python slice_tender.py <docx文件路径> [输出目录] [切片级别]")
        print("\n切片级别说明:")
        print("  0     - 不按章节切片，整个文档转换为一个 Markdown 文件")
        print("  空白  - 按所有标题层级切片（默认）")
        print("  1     - 只按一级标题切片")
        print("  2     - 只按二级标题切片")
        print("  ...   - 以此类推")
        print("\n示例:")
        print("  python slice_tender.py 投标文件.docx")
        print("  python slice_tender.py 投标文件.docx sliced_output")
        print("  python slice_tender.py 投标文件.docx sliced_output 0  # 零级模式")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "sliced"

    # 解析切片级别
    slice_level = None
    if len(sys.argv) > 3:
        try:
            slice_level = int(sys.argv[3])
        except ValueError:
            print(f"错误: 切片级别必须是数字，得到 '{sys.argv[3]}'")
            sys.exit(1)

    try:
        slicer = TenderSlicer(docx_path, output_dir, slice_level)
        slicer.process()
    except Exception as e:
        print(f"错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
