#!/usr/bin/env python3
"""
标书切片工具 - 图形界面版
"""

import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
from pathlib import Path
import threading
import sys
from datetime import datetime
import re

try:
    from docx import Document
except ImportError:
    print("错误: 请先安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


class TenderSlicerApp:
    """标书切片 GUI 应用"""

    def __init__(self, root):
        self.root = root
        self.root.title("标书切片工具")
        self.root.geometry("600x500")
        self.root.resizable(True, True)

        self.input_file = None
        self.output_dir = None

        self.create_widgets()

    def create_widgets(self):
        """创建界面组件"""

        # 标题
        title_frame = tk.Frame(self.root, pady=10)
        title_frame.pack(fill=tk.X)
        tk.Label(title_frame, text="标书切片工具", font=("微软雅黑", 18, "bold")).pack()

        # 文件选择区域
        file_frame = tk.LabelFrame(self.root, text="选择投标文件", padx=10, pady=10)
        file_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Button(file_frame, text="📁 选择文件", command=self.select_file,
                 font=("微软雅黑", 10), bg="#4A90E2", fg="white", width=15).pack(side=tk.LEFT)

        self.file_label = tk.Label(file_frame, text="未选择文件", font=("微软雅黑", 9), fg="#666")
        self.file_label.pack(side=tk.LEFT, padx=10)

        # 输出目录区域
        output_frame = tk.LabelFrame(self.root, text="输出目录", padx=10, pady=10)
        output_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Button(output_frame, text="📂 选择目录", command=self.select_output_dir,
                 font=("微软雅黑", 10), bg="#4A90E2", fg="white", width=15).pack(side=tk.LEFT)

        self.output_label = tk.Label(output_frame, text="默认: ./sliced", font=("微软雅黑", 9), fg="#666")
        self.output_label.pack(side=tk.LEFT, padx=10)

        # 操作按钮
        button_frame = tk.Frame(self.root, pady=15)
        button_frame.pack(fill=tk.X)

        tk.Button(button_frame, text="开始切片", command=self.start_slicing,
                 font=("微软雅黑", 12, "bold"), bg="#52C41A", fg="white",
                 width=20, height=2).pack()

        # 进度条
        self.progress_frame = tk.Frame(self.root)
        self.progress_frame.pack(fill=tk.X, padx=10, pady=5)
        self.progress_label = tk.Label(self.progress_frame, text="", font=("微软雅黑", 9))
        self.progress_label.pack()

        # 日志区域
        log_frame = tk.LabelFrame(self.root, text="处理日志", padx=10, pady=10)
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        self.log_text = scrolledtext.ScrolledText(log_frame, font=("微软雅黑", 9),
                                                  height=10, wrap=tk.WORD)
        self.log_text.pack(fill=tk.BOTH, expand=True)

        # 状态栏
        self.status_bar = tk.Label(self.root, text="就绪", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(fill=tk.X)

    def log(self, message):
        """添加日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update()

    def select_file(self):
        """选择文件"""
        file_path = filedialog.askopenfilename(
            title="选择投标文件",
            filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.input_file = Path(file_path)
            self.file_label.config(text=f"✓ {self.input_file.name}")
            self.log(f"已选择文件: {self.input_file.name}")
            self.status_bar.config(text=f"文件: {self.input_file.name}")

    def select_output_dir(self):
        """选择输出目录"""
        dir_path = filedialog.askdirectory(title="选择输出目录")
        if dir_path:
            self.output_dir = Path(dir_path)
            self.output_label.config(text=f"✓ {self.output_dir}")
            self.log(f"输出目录: {self.output_dir}")

    def get_heading_level(self, paragraph):
        """获取标题级别"""
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            match = re.search(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
            if match:
                return int(match.group(1))
        if hasattr(paragraph, '_element'):
            p = paragraph._element
            if p.pPr and p.pPr.outlineLvl:
                return int(p.pPr.outlineLvl.val) + 1
        return 0

    def table_to_markdown(self, table, start_no=1):
        """表格转 Markdown，带编号"""
        if not table.rows:
            return None, start_no

        lines = []
        no = start_no

        # 表头
        header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        lines.append(f"<!-- {no} --> | " + " | ".join(header_cells) + " |")
        lines.append("<!-- " + str(no + 1) + " --> |" + "|".join(["---"] * len(header_cells)) + "|")
        no += 2

        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # 跳过空行
            if all(cell in ("", " ") for cell in cells):
                continue
            lines.append(f"<!-- {no} --> | " + " | ".join(cells) + " |")
            no += 1

        return '\n'.join(lines) + '\n\n', no

    def sanitize_filename(self, filename):
        """清理文件名"""
        import re
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        if len(filename) > 100:
            filename = filename[:100]
        return filename.strip()

    def slice_document(self):
        """切片文档（在后台线程运行）"""
        try:
            self.log("正在读取文档...")
            doc = Document(str(self.input_file))

            # 初始化切片结果
            sections = []
            current_section = {
                'level': 0,
                'title': '封面',
                'content': [],
                'index': 0
            }
            section_index = 0

            line_no = 1  # 全局行号

            # 处理段落
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                level = self.get_heading_level(paragraph)
                text = paragraph.text.strip()

                if not text:
                    continue

                # 跳过目录
                if any(kw in text for kw in ['目录', '目  录', 'CONTENTS']):
                    continue

                if level > 0:
                    if current_section['content']:
                        sections.append(current_section)
                        section_index += 1
                    current_section = {
                        'level': level,
                        'title': text,
                        'content': [],
                        'index': section_index
                    }
                    # 标题也要编号
                    current_section['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                    line_no += 1
                else:
                    current_section['content'].append(f"<!-- {line_no} --> {text}\n")
                    line_no += 1

                # 更新进度
                if i % 50 == 0:
                    progress = int((i / total_paragraphs) * 100)
                    self.root.after(0, lambda p=progress: self.progress_label.config(text=f"解析中... {p}%"))

            # 处理表格（使用带编号的版本）
            self.log("正在处理表格...")
            for table in doc.tables:
                table_md, line_no = self.table_to_markdown(table, line_no)
                if table_md and current_section['content']:
                    current_section['content'].append(table_md)
                    line_no = line_no  # 更新行号

            # 保存最后一个章节
            if current_section['content']:
                sections.append(current_section)

            # 创建输出目录
            if self.output_dir:
                out_dir = self.output_dir
            else:
                out_dir = Path("sliced")
            out_dir.mkdir(exist_ok=True)

            # 保存文件
            self.log(f"开始保存 {len(sections)} 个章节...")
            for section in sections:
                index = section['index'] + 1
                index_str = str(index).zfill(3)
                safe_title = self.sanitize_filename(section['title'])
                filename = f"{index_str}_{safe_title}.md"
                filepath = out_dir / filename

                content = ''.join(section['content'])
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)

            # 生成索引
            index_file = out_dir / "00_index.md"
            with open(index_file, 'w', encoding='utf-8') as f:
                f.write("# 标书切片索引\n\n")
                f.write(f"原文件: {self.input_file.name}\n")
                f.write(f"切片时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"总章节数: {len(sections)}\n\n---\n\n## 章节列表\n\n")
                for section in sections:
                    idx = section['index'] + 1
                    level = section['level']
                    title = section['title']
                    filename = f"{str(idx).zfill(3)}_{self.sanitize_filename(title)}.md"
                    indent = "  " * (level - 1)
                    f.write(f"{indent}- [{idx}. {title}]({filename})\n")

            self.log(f"✓ 完成！已保存 {len(sections)} 个文件")
            self.log(f"输出目录: {out_dir}")
            self.progress_label.config(text="切片完成！")

            self.root.after(0, lambda: messagebox.showinfo("完成", f"切片完成！\n共 {len(sections)} 个章节\n保存到: {out_dir}"))
            self.root.after(0, lambda: self.status_bar.config(text="完成"))

        except Exception as e:
            error_msg = f"错误: {str(e)}"
            self.log(error_msg)
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
            self.root.after(0, lambda: self.status_bar.config(text="出错"))
        finally:
            self.root.after(0, lambda: self.processing_complete())

    def processing_complete(self):
        """处理完成回调"""
        pass

    def start_slicing(self):
        """开始切片"""
        if not self.input_file:
            messagebox.showwarning("提示", "请先选择投标文件")
            return

        # 清空日志
        self.log_text.delete(1.0, tk.END)
        self.log("开始处理...")
        self.progress_label.config(text="初始化中...")
        self.status_bar.config(text="处理中...")

        # 在后台线程执行
        thread = threading.Thread(target=self.slice_document)
        thread.daemon = True
        thread.start()


def main():
    """主函数"""
    root = tk.Tk()

    # 尝试设置图标（Windows）
    try:
        # 可以设置应用图标
        pass
    except:
        pass

    app = TenderSlicerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
